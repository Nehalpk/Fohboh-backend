import concurrent.futures
from concurrent.futures import ThreadPoolExecutor
import threading
import os
import json
import openai
import requests
import numpy as np
import boto3
import psycopg2
from psycopg2.extras import execute_values
from psycopg2 import pool
import pandas as pd
import io
import logging
import time
from fastapi import APIRouter, UploadFile, File, Form, HTTPException, Depends
from fastapi.responses import JSONResponse
from typing import List, Dict, Any, Optional
from pydantic import BaseModel
from dotenv import load_dotenv

import voyageai
import pdfplumber
from docx import Document
import openpyxl

import multiprocessing as mp

# Import the necessary functions from chat_gpt.py for authentication
from src.chat_gpt import get_current_user, get_db
from src.File_upload import verify_restaurant_access

# Load environment variables
load_dotenv()

CPU_COUNT = max(1, mp.cpu_count() - 1)
# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Initialize API clients with keys from environment
OPENAI_API_KEY = "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"
ANTHROPIC_API_KEY = "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"

openai_client = openai.OpenAI(api_key=OPENAI_API_KEY)

# Initialize thread-safe resources
thread_pool = ThreadPoolExecutor(max_workers=5)
openai_embedding_lock = threading.Lock()
claude_embedding_lock = threading.Lock()

# S3 Configuration
BUCKET_NAME = "my-audio-demo"
UPLOAD_BASE_DIR = "uploads/restaurants"

# Initialize S3 client
s3_client = boto3.client(
    's3',
    region_name='us-east-1',
    aws_access_key_id='AKIA6G75DKGISWQMC7CR',
    aws_secret_access_key='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'
)

# Database connection pool with hardcoded values (as provided)
DB_POOL = pool.SimpleConnectionPool(
    1, 10,
    host="database-1-snapshot-26-july.cboosuomg0xi.us-east-1.rds.amazonaws.com",
    database="postgres",
    user="postgres",
    password="EqF3XKOz13DX3jE6APrW"
)

# Create router
router = APIRouter(prefix="/rag", tags=["RAG"])

# Valid categories (same as File_upload.py)
VALID_CATEGORIES = ["Inventory", "Labor", "Sales", "Menu"]


# Models
class DeleteEmbeddingsRequest(BaseModel):
    restaurant_name: str
    category: str
    file_names: List[str]


class QueryEmbeddingsRequest(BaseModel):
    restaurant_name: str
    category: str
    query: str
    top_k: int = 5
    embedding_model: str = "openai"  # or "claude"


# Helper Functions
def get_restaurant_folder_path(restaurant_name: str) -> str:
    """Create sanitized folder path for a restaurant"""
    safe_name = restaurant_name.replace(" ", "_").lower()
    return f"{UPLOAD_BASE_DIR}/{safe_name}"


def get_user_restaurant_path(restaurant_name: str, user_id: int, category: str) -> str:
    """Get the S3 path for a user's restaurant category folder"""
    # if category not in VALID_CATEGORIES:
    #     raise ValueError(f"Invalid category. Must be one of {VALID_CATEGORIES}")

    restaurant_folder = get_restaurant_folder_path(restaurant_name)
    return f"{restaurant_folder}/user_{user_id}/{category}"


def is_valid_integer(value):
    """Check if a value can be converted to an integer"""
    try:
        int(value)
        return True
    except (ValueError, TypeError):
        return False


def init_embedding_tables():
    """Initialize database tables for embeddings if they don't exist"""
    conn = None
    try:
        conn = DB_POOL.getconn()
        cur = conn.cursor()

        # Make sure pgvector extension is enabled
        cur.execute("CREATE EXTENSION IF NOT EXISTS vector")

        # Create OpenAI embeddings table with embedding as vector (1536 dimensions)
        cur.execute("""
        CREATE TABLE IF NOT EXISTS openai_embeddings (
            id SERIAL PRIMARY KEY,
            user_id INTEGER NOT NULL,
            restaurant_name TEXT NOT NULL,
            category TEXT NOT NULL,
            filename TEXT NOT NULL,
            chunk_id INTEGER NOT NULL,
            text_content TEXT NOT NULL,
            embedding VECTOR(1536) NOT NULL,
            created_at TIMESTAMP WITH TIME ZONE DEFAULT CURRENT_TIMESTAMP
        )
        """)

        # Create Claude embeddings table with embedding as vector (1024 dimensions)
        cur.execute("""
        CREATE TABLE IF NOT EXISTS claude_embeddings (
            id SERIAL PRIMARY KEY,
            user_id INTEGER NOT NULL,
            restaurant_name TEXT NOT NULL,
            category TEXT NOT NULL,
            filename TEXT NOT NULL,
            chunk_id INTEGER NOT NULL,
            text_content TEXT NOT NULL,
            embedding VECTOR(1024) NOT NULL,
            created_at TIMESTAMP WITH TIME ZONE DEFAULT CURRENT_TIMESTAMP
        )
        """)

        # Create chat_history embeddings table with embedding as vector (1536 dimensions)
        cur.execute("""
        CREATE TABLE IF NOT EXISTS chat_history_claude (
            id SERIAL PRIMARY KEY,

            user_id INTEGER NOT NULL REFERENCES managers(id) ON DELETE CASCADE,

            restaurant_names TEXT[],  -- list of restaurant names (array of text)

            conversation_id VARCHAR(100) NOT NULL,  -- logical thread
            message_id VARCHAR(100) NOT NULL,       -- unique message identifier

            model VARCHAR(50) DEFAULT 'claude',                      -- e.g., 'claude-3-sonnet'

            timestamp TIMESTAMPTZ DEFAULT CURRENT_TIMESTAMP,  -- with timezone support

            question TEXT NOT NULL,
            answer TEXT NOT NULL,

            -- Optional columns for future extensibility
            source TEXT,                            -- source of question (e.g., UI, API)
            tags TEXT[],                            -- for future categorization/search

            UNIQUE(conversation_id, message_id)     -- prevent duplicate messages per convo
        )

        """)

        # Create chat_history embeddings table with embedding as vector (1536 dimensions)
        cur.execute("""
        CREATE TABLE IF NOT EXISTS chat_history_openai (
            id SERIAL PRIMARY KEY,

            user_id INTEGER NOT NULL REFERENCES managers(id) ON DELETE CASCADE,

            restaurant_names TEXT[],  -- list of restaurant names (array of text)

            conversation_id VARCHAR(100) NOT NULL,  -- logical thread
            message_id VARCHAR(100) NOT NULL,       -- unique message identifier

            model VARCHAR(50) DEFAULT 'openai',           

            timestamp TIMESTAMPTZ DEFAULT CURRENT_TIMESTAMP,  -- with timezone support

            question TEXT NOT NULL,
            answer TEXT NOT NULL,

            -- Optional columns for future extensibility
            source TEXT,                            -- source of question (e.g., UI, API)
            tags TEXT[],                            -- for future categorization/search

            UNIQUE(conversation_id, message_id)     -- prevent duplicate messages per convo
        )
        """)

        # Create sales_graphs table with only the essential columns for graphs
        # SALES GRAPHS
        # cur.execute("""
        # CREATE TABLE IF NOT EXISTS sales_graphs (
        #     id SERIAL PRIMARY KEY,
        #     restaurant_id INTEGER NOT NULL,
        #     sale_id VARCHAR(100) NOT NULL,
        #     date DATE NOT NULL,
        #     time TIME NOT NULL,
        #     items_sold TEXT NOT NULL,
        #     number_of_items INTEGER NOT NULL,
        #     subtotal DECIMAL(10, 2) NOT NULL,
        #     tip DECIMAL(10, 2),
        #     total_amount DECIMAL(10, 2) NOT NULL,
        #     payment_method VARCHAR(50) NOT NULL,
        #     order_type VARCHAR(50) NOT NULL,
        #     filename TEXT NOT NULL,
        #     created_at TIMESTAMP WITH TIME ZONE DEFAULT CURRENT_TIMESTAMP,
        #     FOREIGN KEY (restaurant_id) REFERENCES restaurants(id)
        # )
        # """)

        # cur.execute("""
        # CREATE INDEX IF NOT EXISTS sales_graphs_restaurant_id_idx
        # ON sales_graphs(restaurant_id)
        # """)

        # cur.execute("""
        # CREATE INDEX IF NOT EXISTS sales_graphs_date_idx
        # ON sales_graphs(date)
        # """)

        # cur.execute("""
        # CREATE INDEX IF NOT EXISTS sales_graphs_restaurant_date_idx
        # ON sales_graphs(restaurant_id, date)
        # """)

        # # INVENTORY GRAPHS
        # cur.execute("""
        # CREATE TABLE IF NOT EXISTS inventory_graphs (
        #     id SERIAL PRIMARY KEY,
        #     restaurant_id INTEGER NOT NULL,
        #     date DATE NOT NULL,
        #     ingredient TEXT NOT NULL,
        #     quantity DECIMAL(10, 2) NOT NULL,
        #     par_level DECIMAL(10, 2),
        #     unit_cost DECIMAL(10, 2),
        #     is_low BOOLEAN,
        #     filename TEXT NOT NULL,
        #     created_at TIMESTAMP WITH TIME ZONE DEFAULT CURRENT_TIMESTAMP,
        #     FOREIGN KEY (restaurant_id) REFERENCES restaurants(id)
        # )
        # """)

        # cur.execute("""
        # CREATE INDEX IF NOT EXISTS inventory_graphs_restaurant_id_idx
        # ON inventory_graphs(restaurant_id)
        # """)

        # cur.execute("""
        # CREATE INDEX IF NOT EXISTS inventory_graphs_date_idx
        # ON inventory_graphs(date)
        # """)

        # # MENU GRAPHS
        # cur.execute("""
        # CREATE TABLE IF NOT EXISTS menu_graphs (
        #     id SERIAL PRIMARY KEY,
        #     restaurant_id INTEGER NOT NULL,
        #     menu_item TEXT NOT NULL,
        #     ingredient TEXT NOT NULL,
        #     amount DECIMAL(10, 2),
        #     unit_cost DECIMAL(10, 2),
        #     filename TEXT NOT NULL,
        #     created_at TIMESTAMP WITH TIME ZONE DEFAULT CURRENT_TIMESTAMP,
        #     FOREIGN KEY (restaurant_id) REFERENCES restaurants(id)
        # )
        # """)

        # cur.execute("""
        # CREATE INDEX IF NOT EXISTS menu_graphs_restaurant_id_idx
        # ON menu_graphs(restaurant_id)
        # """)

        # # EMPLOYEE GRAPHS
        # cur.execute("""
        # CREATE TABLE IF NOT EXISTS employee_graphs (
        #     id SERIAL PRIMARY KEY,
        #     restaurant_id INTEGER NOT NULL,
        #     employee_id INTEGER NOT NULL,
        #     name TEXT NOT NULL,
        #     role TEXT NOT NULL,
        #     hire_date DATE NOT NULL,
        #     termination_date DATE,
        #     hourly_rate DECIMAL(10, 2),
        #     filename TEXT NOT NULL,
        #     created_at TIMESTAMP WITH TIME ZONE DEFAULT CURRENT_TIMESTAMP,
        #     FOREIGN KEY (restaurant_id) REFERENCES restaurants(id)
        # )
        # """)

        # cur.execute("""
        # CREATE INDEX IF NOT EXISTS employee_graphs_restaurant_id_idx
        # ON employee_graphs(restaurant_id)
        # """)

        # # PROMOTION GRAPHS
        # cur.execute("""
        # CREATE TABLE IF NOT EXISTS promotion_graphs (
        #     id SERIAL PRIMARY KEY,
        #     restaurant_id INTEGER NOT NULL,
        #     promotion_id INTEGER NOT NULL,
        #     promotion_name TEXT NOT NULL,
        #     start_date DATE,
        #     end_date DATE,
        #     discount TEXT,
        #     filename TEXT NOT NULL,
        #     created_at TIMESTAMP WITH TIME ZONE DEFAULT CURRENT_TIMESTAMP,
        #     FOREIGN KEY (restaurant_id) REFERENCES restaurants(id)
        # )
        # """)

        # Create indexes for faster similarity search
        cur.execute("""
        CREATE INDEX IF NOT EXISTS openai_embedding_idx ON openai_embeddings 
        USING ivfflat (embedding vector_l2_ops) WITH (lists = 100)
        """)

        cur.execute("""
        CREATE INDEX IF NOT EXISTS claude_embedding_idx ON claude_embeddings 
        USING ivfflat (embedding vector_l2_ops) WITH (lists = 100)
        """)

        # Create compound indexes for faster filtering
        cur.execute("""
        CREATE INDEX IF NOT EXISTS openai_user_restaurant_category_idx 
        ON openai_embeddings(user_id, restaurant_name, category)
        """)

        cur.execute("""
        CREATE INDEX IF NOT EXISTS claude_user_restaurant_category_idx 
        ON claude_embeddings(user_id, restaurant_name, category)
        """)

        # Create indexes for faster similarity search
        # cur.execute("""
        # CREATE INDEX IF NOT EXISTS chat_history_claude_idx ON chat_history_claude
        # USING ivfflat (embedding vector_l2_ops) WITH (lists = 100)
        # """)

        # # Create indexes for faster similarity search
        # cur.execute("""
        # CREATE INDEX IF NOT EXISTS chat_history_openai_idx ON chat_history_openai
        # USING ivfflat (embedding vector_l2_ops) WITH (lists = 100)
        # """)

        conn.commit()
        logger.info("Embedding tables initialized successfully")

    except Exception as e:
        if conn:
            conn.rollback()
        logger.error(f"Error initializing embedding tables: {str(e)}")
        raise
    finally:
        if conn:
            cur.close()
            DB_POOL.putconn(conn)


async def check_embedding_exit(user_id: int, category: str, filename: str, embedding_model: str,
                               restaurant_name: str) -> bool:
    """
    Checking if the embedding already exist for a specific file
    """
    try:
        table = "openai_embeddings" if embedding_model == "openai" else "claude_embeddings"
        conn = DB_POOL.getconn()
        try:
            cur = conn.cursor()
            cur.execute(f"""
                SELECT 
                    COUNT(id) as count
                FROM {table}
                WHERE 
                    user_id = %s AND
                    restaurant_name = %s AND
                    category = %s AND
                    filename = %s
            """, (user_id, restaurant_name, category, filename))
            result = cur.fetchone()

            # Check if result exists and has a count greater than 0
            return result is not None and result[0] > 0

        finally:
            cur.close()
            DB_POOL.putconn(conn)

    except Exception as e:
        logger.error(f"Error in check_embedding_exit: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Error in check_embedding_exit: {str(e)}"
        )


def process_chunk_openai(text, user_id, restaurant_name, category, filename, chunk_id):
    """Process a text chunk and generate OpenAI embeddings"""
    try:
        with openai_embedding_lock:
            response = openai_client.embeddings.create(
                input=[text],
                model="text-embedding-ada-002"
            )
            # Convert numpy values to Python float
            embedding = [float(x) for x in response.data[0].embedding]
        return (user_id, restaurant_name, category, filename, chunk_id, text, embedding)
    except Exception as e:
        logger.error(f"Error processing OpenAI chunk {chunk_id}: {str(e)}")
        return None


def get_claude_embedding(text):
    """Get embeddings from Claude's voyage-multimodal-3 model"""
    try:

        vo = voyageai.Client(api_key="pa-Es-vdxX3QRYunzmZnD_gTcnZ4uvdT1nmR6vykvbNDux")

        result = vo.embed(["hello world"], model="voyage-3-large")

        embedding = result.embeddings[0]

        # Ensure we have 1024 dimensions
        if len(embedding) != 1024:
            logger.warning(f"Expected 1024 dimensions, got {len(embedding)}")

        # Convert to Python floats
        return [float(x) for x in embedding]
    except Exception as e:
        logger.error(f"Error getting Claude embedding: {str(e)}")
        raise e


def process_chunk_claude(text, user_id, restaurant_name, category, filename, chunk_id):
    """Process a text chunk and create Claude embeddings using voyage-multimodal-3"""
    try:
        with claude_embedding_lock:
            embedding = get_claude_embedding(text)
            return (user_id, restaurant_name, category, filename, chunk_id, text, embedding)
    except Exception as e:
        logger.error(f"Error processing Claude chunk {chunk_id}: {str(e)}")
        return None


# async def verify_restaurant_access(restaurant_name: str, current_user: dict, conn) -> dict:
#     """Verify user has access to the restaurant and it exists"""
#     try:
#         cur = conn.cursor()

#         # Different query based on user role
#         if current_user["role"] == "SUPER_ADMIN":
#             # SUPER_ADMIN can access any active restaurant
#             cur.execute("""
#                 SELECT id, name, location
#                 FROM restaurants
#                 WHERE name = %s AND active = true
#             """, (restaurant_name,))
#         else:
#             # Regional and Restaurant managers can only access assigned restaurants
#             cur.execute("""
#                 SELECT r.id, r.name, r.location
#                 FROM restaurants r
#                 JOIN restaurant_assignments ra ON r.id = ra.restaurant_id
#                 WHERE r.name = %s AND r.active = true AND ra.manager_id = %s
#             """, (restaurant_name, current_user["id"]))

#         restaurant = cur.fetchone()
#         if not restaurant:
#             raise HTTPException(
#                 status_code=404,
#                 detail=f"Restaurant '{restaurant_name}' not found or you don't have access to it"
#             )

#         return dict(restaurant)

#     except HTTPException:
#         raise
#     except Exception as e:
#         logger.error(f"Error verifying restaurant access: {str(e)}")
#         raise HTTPException(
#             status_code=500,
#             detail=f"Error verifying restaurant access: {str(e)}"
#         )

# Main functions for embedding generation and retrieval
async def generate_embeddings_for_csv(
        s3_key: str,
        restaurant_name: str,
        category: str,
        filename: str,
        embedding_model: str,
        current_user: dict,
        batch_size: int = 10
) -> Dict[str, Any]:
    """Generate embeddings for a CSV file stored in S3"""
    try:
        # Retrieve file from S3
        response = s3_client.get_object(
            Bucket=BUCKET_NAME,
            Key=s3_key
        )

        file_content = response['Body'].read()
        df = pd.read_csv(io.StringIO(file_content.decode('utf-8')))

        if 'Record Number' in df.columns:
            df = df[df['Record Number'] != 'Record Number']

        # Create chunks (one per row)
        chunks = []
        for index, row in df.iterrows():
            text = " ".join([f"{col}: {row[col]}" for col in df.columns])
            chunk_id = index if is_valid_integer(index) else hash(text) % (2 ** 31)

            if embedding_model == "openai":
                chunks.append((text, current_user["id"], restaurant_name, category, filename, chunk_id))
            else:  # claude
                chunks.append((text, current_user["id"], restaurant_name, category, filename, chunk_id))

        # Process chunks in batches
        all_embeddings = []

        # Choose the appropriate processing function and table
        if embedding_model == "openai":
            process_func = process_chunk_openai
            table_name = "openai_embeddings"
        else:
            process_func = process_chunk_claude
            table_name = "claude_embeddings"

        # Process chunks in parallel
        with ThreadPoolExecutor(max_workers=5) as executor:
            futures = [executor.submit(process_func, *chunk) for chunk in chunks]
            for future in concurrent.futures.as_completed(futures):
                result = future.result()
                if result:
                    all_embeddings.append(result)

                # Insert in batches when we have enough embeddings
                if len(all_embeddings) >= batch_size:
                    conn = DB_POOL.getconn()
                    try:
                        cur = conn.cursor()
                        execute_values(cur,
                                       f"INSERT INTO {table_name} (user_id, restaurant_name, category, filename, chunk_id, text_content, embedding) VALUES %s",
                                       all_embeddings
                                       )
                        conn.commit()
                        all_embeddings = []
                    finally:
                        cur.close()
                        DB_POOL.putconn(conn)

        # Insert any remaining embeddings
        if all_embeddings:
            conn = DB_POOL.getconn()
            try:
                cur = conn.cursor()
                execute_values(cur,
                               f"INSERT INTO {table_name} (user_id, restaurant_name, category, filename, chunk_id, text_content, embedding) VALUES %s",
                               all_embeddings
                               )
                conn.commit()
            finally:
                cur.close()
                DB_POOL.putconn(conn)

        return {
            "message": f"{embedding_model.capitalize()} embeddings for {filename} processed and stored successfully.",
            "restaurant": restaurant_name,
            "category": category,
            "rows_processed": len(df),
            "embedding_model": embedding_model
        }

    except Exception as e:
        logger.error(f"Error generating embeddings: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


# API Endpoints
@router.post("/embeddings/generate-openai")
async def generate_openai_embeddings(
        restaurant_name: str = Form(...),
        category: str = Form(...),
        filename: str = Form(...),
        current_user: dict = Depends(get_current_user),
        conn=Depends(get_db)
):
    """Generate OpenAI embeddings for file"""
    try:
        # Verify restaurant access
        restaurant = await verify_restaurant_access(restaurant_name, current_user)

        # Validate category
        # if category not in VALID_CATEGORIES:
        #     raise HTTPException(
        #         status_code=400,
        #         detail=f"Category must be one of: {', '.join(VALID_CATEGORIES)}"
        #     )

        # Check if the embedding already exist
        embeddings_exit = await check_embedding_exit(
            restaurant_name=restaurant_name,
            category=category,
            filename=filename,
            embedding_model="openai",
            user_id=current_user["id"]
        )

        if embeddings_exit:
            return {
                "message": f"Embeddings for {filename} already exist",
                "restaurant": restaurant_name,
                "category": category,
                "filename": filename,
                "statue": "already_embedded"
            }
        # Construct S3 path
        s3_folder = get_user_restaurant_path(restaurant_name, current_user["id"], category)
        s3_key = f"{s3_folder}/{filename}"

        # Check if file exists
        try:
            s3_client.head_object(
                Bucket=BUCKET_NAME,
                Key=s3_key
            )
        except Exception:
            raise HTTPException(
                status_code=404,
                detail=f"File '{filename}' not found in {restaurant_name}/{category}"
            )

        # Generate embeddings
        # result = await generate_embeddings_for_csv(

        result = await generate_embeddings_all_docs(
            s3_key=s3_key,
            restaurant_name=restaurant_name,
            category=category,
            filename=filename,
            embedding_model="openai",
            current_user=current_user
        )

        return result

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error in OpenAI embedding generation: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Error generating OpenAI embeddings: {str(e)}"
        )


@router.post("/embeddings/generate-claude")
async def generate_claude_embeddings(
        restaurant_name: str = Form(...),
        category: str = Form(...),
        filename: str = Form(...),
        current_user: dict = Depends(get_current_user),
        conn=Depends(get_db)
):
    """Generate Claude embeddings for a specific CSV file using voyage-multimodal-3"""
    try:
        # Verify restaurant access
        restaurant = await verify_restaurant_access(restaurant_name, current_user)

        # Validate category
        # if category not in VALID_CATEGORIES:
        #     raise HTTPException(
        #         status_code=400,
        #         detail=f"Category must be one of: {', '.join(VALID_CATEGORIES)}"
        #     )

        embeddings_exit = await check_embedding_exit(
            restaurant_name=restaurant_name,
            category=category,
            filename=filename,
            embedding_model="claude",
            user_id=current_user["id"]
        )

        if embeddings_exit:
            return {
                "message": f"Embeddings for {filename} already exist",
                "restaurant": restaurant_name,
                "category": category,
                "filename": filename,
                "statue": "already_embedded"
            }

        # Construct S3 path
        s3_folder = get_user_restaurant_path(restaurant_name, current_user["id"], category)
        s3_key = f"{s3_folder}/{filename}"

        # Check if file exists
        try:
            s3_client.head_object(
                Bucket=BUCKET_NAME,
                Key=s3_key
            )
        except Exception:
            raise HTTPException(
                status_code=404,
                detail=f"File '{filename}' not found in {restaurant_name}/{category}"
            )

        # Generate embeddings
        # result = await generate_embeddings_for_csv(

        result = await generate_embeddings_all_docs(
            s3_key=s3_key,
            restaurant_name=restaurant_name,
            category=category,
            filename=filename,
            embedding_model="claude",
            current_user=current_user
        )

        return result

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error in Claude embedding generation: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Error generating Claude embeddings: {str(e)}"
        )


@router.get("/embeddings/list")
async def list_embedding_files(
        restaurant_name: str,
        # category: str,
        embedding_model: str = "claude",  # or "openai"
        current_user: dict = Depends(get_current_user),
        conn=Depends(get_db)
):
    """List files that have embeddings for a specific restaurant and category"""
    try:
        # Verify restaurant access
        restaurant = await verify_restaurant_access(restaurant_name, current_user)

        # Validate embedding model
        if embedding_model not in ["openai", "claude"]:
            raise HTTPException(
                status_code=400,
                detail="Embedding model must be 'openai' or 'claude'"
            )

        # Get the appropriate table name
        table = "openai_embeddings" if embedding_model == "openai" else "claude_embeddings"

        # Query for files
        conn = DB_POOL.getconn()
        try:
            cur = conn.cursor()

            # Log for debugging
            # logger.info(f"Listing embeddings for restaurant: {restaurant_name}, user: {current_user['id']}, role: {current_user['role']}")

            # For all user roles, once they have access to the restaurant,
            # they should see all embeddings for that restaurant

            # First, let's check if there are any embeddings for this restaurant
            cur.execute(f"""
                SELECT COUNT(*) as count
                FROM {table}
                WHERE restaurant_name = %s
            """, (restaurant_name,))

            count_result = cur.fetchone()
            logger.info(f"Total embeddings for restaurant '{restaurant_name}': {count_result[0]}")

            # If no embeddings found with exact match, try case-insensitive match
            if count_result[0] == 0:
                logger.info(f"No embeddings found with exact match, trying case-insensitive match")

                # List all restaurant names in the embeddings table for debugging
                cur.execute(f"""
                    SELECT DISTINCT restaurant_name
                    FROM {table}
                    ORDER BY restaurant_name
                """)
                all_restaurants = cur.fetchall()
                logger.info(f"All restaurant names in {table}: {[row[0] for row in all_restaurants]}")

                # Try case-insensitive match
                cur.execute(f"""
                    SELECT DISTINCT restaurant_name
                    FROM {table}
                    WHERE LOWER(restaurant_name) = LOWER(%s)
                """, (restaurant_name,))

                similar_names = cur.fetchall()
                if similar_names:
                    logger.info(f"Found similar restaurant names: {[row[0] for row in similar_names]}")
                    # Use the first similar name found
                    restaurant_name = similar_names[0][0]
                    logger.info(f"Using restaurant name: {restaurant_name}")

            # Now get the actual files
            cur.execute(f"""
                SELECT 
                    filename,
                    category,
                    COUNT(DISTINCT chunk_id) as chunk_count,
                    MAX(created_at) as last_updated,
                    user_id
                FROM {table}
                WHERE 
                    restaurant_name = %s
                GROUP BY filename, category, user_id
                ORDER BY last_updated DESC
            """, (restaurant_name,))

            rows = cur.fetchall()
            logger.info(f"Found {len(rows)} embedding files for restaurant: {restaurant_name}")

            # Get user information for each file
            files = []
            for row in rows:
                # Get user info for this embedding
                cur.execute("""
                    SELECT 
                        email, 
                        full_name, 
                        role
                    FROM managers 
                    WHERE id = %s
                """, (row[4],))

                user_info = cur.fetchone()
                user_display = "Unknown User"

                if user_info:
                    user_display = user_info[1] if user_info[1] else user_info[0]

                files.append({
                    "filename": row[0],
                    "category": row[1],
                    "chunk_count": row[2],
                    "last_updated": row[3].strftime("%Y-%m-%d %H:%M:%S") if row[3] else None,
                    "created_by": user_display,
                    "creator_role": user_info[2] if user_info else "Unknown"
                })

            return {
                "restaurant": restaurant_name,
                "embedding_model": embedding_model,
                "files": files,
                "count": len(files)
            }

        finally:
            cur.close()
            DB_POOL.putconn(conn)

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error listing embedding files: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Error listing embedding files: {str(e)}"
        )


@router.get("/embeddings/list-all")
async def list_all_embeddings(
        restaurant_name: str,
        current_user: dict = Depends(get_current_user),
        conn=Depends(get_db)
):
    """List all embeddings for a specific restaurant across both OpenAI and Claude models"""
    try:
        # Verify restaurant access
        restaurant = await verify_restaurant_access(restaurant_name, current_user)

        # Define the tables for both models
        tables = {
            "openai": "openai_embeddings",
            "claude": "claude_embeddings"
        }

        all_files = []

        # Query for each model
        conn = DB_POOL.getconn()
        try:
            cur = conn.cursor()

            for model, table in tables.items():
                # Check if there are any embeddings for this restaurant
                cur.execute(f"""
                    SELECT COUNT(*) as count
                    FROM {table}
                    WHERE restaurant_name = %s
                """, (restaurant_name,))
                count_result = cur.fetchone()

                if count_result[0] == 0:
                    logger.info(f"No embeddings found for restaurant '{restaurant_name}' in {model} model")
                    continue

                # Fetch embedding files
                cur.execute(f"""
                    SELECT 
                        filename,
                        category,
                        COUNT(DISTINCT chunk_id) as chunk_count,
                        MAX(created_at) as last_updated,
                        user_id
                    FROM {table}
                    WHERE 
                        restaurant_name = %s
                    GROUP BY filename, category, user_id
                    ORDER BY last_updated DESC
                """, (restaurant_name,))

                rows = cur.fetchall()

                for row in rows:
                    # Get user info for this embedding
                    cur.execute("""
                        SELECT 
                            email, 
                            full_name, 
                            role
                        FROM managers 
                        WHERE id = %s
                    """, (row[4],))

                    user_info = cur.fetchone()
                    user_display = "Unknown User"

                    if user_info:
                        user_display = user_info[1] if user_info[1] else user_info[0]

                    all_files.append({
                        "filename": row[0],
                        "category": row[1],
                        "chunk_count": row[2],
                        "last_updated": row[3].strftime("%Y-%m-%d %H:%M:%S") if row[3] else None,
                        "created_by": user_display,
                        "creator_role": user_info[2] if user_info else "Unknown",
                        "embedding_model": model
                    })

            return {
                "restaurant": restaurant_name,
                "files": all_files,
                "count": len(all_files)
            }

        finally:
            cur.close()
            DB_POOL.putconn(conn)

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error listing all embeddings: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Error listing all embeddings: {str(e)}"
        )


@router.get("/embeddings/cat_list")
async def list_embedding_cat_files(
        restaurant_name: str,
        category: str,
        embedding_model: str = "openai",  # or "claude"
        current_user: dict = Depends(get_current_user),
        conn=Depends(get_db)
):
    """List files that have embeddings for a specific restaurant and category"""
    try:
        # Verify restaurant access
        restaurant = await verify_restaurant_access(restaurant_name, current_user)

        # Validate category
        # if category not in VALID_CATEGORIES:
        #     raise HTTPException(
        #         status_code=400,
        #         detail=f"Category must be one of: {', '.join(VALID_CATEGORIES)}"
        #     )

        # Validate embedding model
        if embedding_model not in ["openai", "claude"]:
            raise HTTPException(
                status_code=400,
                detail="Embedding model must be 'openai' or 'claude'"
            )

        # Get the appropriate table name
        table = "openai_embeddings" if embedding_model == "openai" else "claude_embeddings"

        # Query for files
        conn = DB_POOL.getconn()
        try:
            cur = conn.cursor()

            cur.execute(f"""
                SELECT 
                    filename,
                    COUNT(DISTINCT chunk_id) as chunk_count,
                    MAX(created_at) as last_updated
                FROM {table}
                WHERE 
                    user_id = %s AND
                    restaurant_name = %s AND
                    category = %s
                GROUP BY filename
                ORDER BY last_updated DESC
            """, (current_user["id"], restaurant_name, category))

            rows = cur.fetchall()

            files = [{
                "filename": row[0],
                "chunk_count": row[1],
                "last_updated": row[2].strftime("%Y-%m-%d %H:%M:%S") if row[2] else None
            } for row in rows]

            return {
                "restaurant": restaurant_name,
                "category": category,
                "embedding_model": embedding_model,
                "files": files,
                "count": len(files)
            }

        finally:
            cur.close()
            DB_POOL.putconn(conn)

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error listing embedding files: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Error listing embedding files: {str(e)}"
        )


@router.delete("/embeddings/delete")
async def delete_embeddings(
        request: DeleteEmbeddingsRequest,
        current_user: dict = Depends(get_current_user),
        conn=Depends(get_db)
):
    """Delete embeddings for specific files"""
    try:
        # Verify restaurant access
        restaurant = await verify_restaurant_access(request.restaurant_name, current_user)

        # Validate category
        # if request.category not in VALID_CATEGORIES:
        #     raise HTTPException(
        #         status_code=400,
        #         detail=f"Category must be one of: {', '.join(VALID_CATEGORIES)}"
        #     )

        # Delete embeddings from both tables
        conn = DB_POOL.getconn()
        try:
            cur = conn.cursor()

            openai_deleted = 0
            claude_deleted = 0

            for filename in request.file_names:
                # Delete from OpenAI embeddings
                cur.execute("""
                    DELETE FROM openai_embeddings
                    WHERE 
                        user_id = %s AND
                        restaurant_name = %s AND
                        category = %s AND
                        filename = %s
                    RETURNING id
                """, (current_user["id"], request.restaurant_name, request.category, filename))

                openai_deleted += cur.rowcount

                # Delete from Claude embeddings
                cur.execute("""
                    DELETE FROM claude_embeddings
                    WHERE 
                        user_id = %s AND
                        restaurant_name = %s AND
                        category = %s AND
                        filename = %s
                    RETURNING id
                """, (current_user["id"], request.restaurant_name, request.category, filename))

                claude_deleted += cur.rowcount

            conn.commit()

            return {
                "message": f"Successfully deleted embeddings for {len(request.file_names)} files",
                "openai_records_deleted": openai_deleted,
                "claude_records_deleted": claude_deleted,
                "restaurant": request.restaurant_name,
                "category": request.category
            }

        except Exception as e:
            conn.rollback()
            logger.error(f"Error deleting embeddings: {str(e)}")
            raise HTTPException(
                status_code=500,
                detail=f"Error deleting embeddings: {str(e)}"
            )
        finally:
            cur.close()
            DB_POOL.putconn(conn)

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error in delete_embeddings: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Error deleting embeddings: {str(e)}"
        )


# @router.post("/query")
# async def query_embeddings(
#     request: QueryEmbeddingsRequest,
#     current_user: dict = Depends(get_current_user),
#     conn = Depends(get_db)
# ):
#     """Query embeddings for similar content, prioritizing chat history first."""
#     try:
#         # Verify restaurant access
#         restaurant = await verify_restaurant_access(request.restaurant_name, current_user, conn)

#         # Validate category
#         if request.category not in VALID_CATEGORIES:
#             raise HTTPException(
#                 status_code=400,
#                 detail=f"Category must be one of: {', '.join(VALID_CATEGORIES)}"
#             )

#         # Validate embedding model
#         if request.embedding_model not in ["openai", "claude"]:
#             raise HTTPException(
#                 status_code=400,
#                 detail="Embedding model must be 'openai' or 'claude'"
#             )

#         # Generate query embedding using the appropriate model
#         if request.embedding_model == "openai":
#             with openai_embedding_lock:
#                 response = openai_client.embeddings.create(
#                     input=[request.query],
#                     model="text-embedding-ada-002"
#                 )
#                 query_embedding = [float(x) for x in response.data[0].embedding]
#         else:  # claude
#             try:
#                 query_embedding = get_claude_embedding(request.query)
#             except Exception as e:
#                 logger.error(f"Error getting Claude embedding, falling back to OpenAI: {str(e)}")
#                 with openai_embedding_lock:
#                     response = openai_client.embeddings.create(
#                         input=[request.query],
#                         model="text-embedding-ada-002"
#                     )
#                     full_embedding = response.data[0].embedding

#                     # Reduce dimensions to 1024
#                     claude_embedding = [
#                         float(sum(full_embedding[i:i+3]) / min(3, len(full_embedding) - i))
#                         for i in range(0, 1536, 3)
#                     ]

#                     # Ensure exactly 1024 dimensions
#                     if len(claude_embedding) < 1024:
#                         claude_embedding.extend([0.0] * (1024 - len(claude_embedding)))

#                     # Normalize embedding
#                     norm = np.linalg.norm(claude_embedding)
#                     query_embedding = [float(x) for x in (claude_embedding / norm)] if norm > 0 else claude_embedding

#         # **Step 1: Check chat history first**
#         history_results = await search_chat_history(request, current_user, conn)

#         # If a highly relevant past chat exists, return it instead of querying embeddings
#         if history_results and history_results[0]["similarity"] > 0.75:
#             return {
#                 "query": request.query,
#                 "restaurant": request.restaurant_name,
#                 "category": request.category,
#                 "embedding_model": request.embedding_model,
#                 "results": history_results,
#                 "count": len(history_results),
#                 "source": "chat_history"
#             }

#         # **Step 2: No good history match, search in embeddings**
#         table_name = "openai_embeddings" if request.embedding_model == "openai" else "claude_embeddings"

#         conn = DB_POOL.getconn()
#         try:
#             cur = conn.cursor()

#             embedding_str = "[" + ",".join(str(x) for x in query_embedding) + "]"

#             cur.execute(f"""
#                 SELECT
#                     id,
#                     filename,
#                     chunk_id,
#                     text_content,
#                     1 - (embedding <=> %s::vector) AS similarity
#                 FROM {table_name}
#                 WHERE
#                     user_id = %s AND
#                     restaurant_name = %s AND
#                     category = %s
#                 ORDER BY embedding <=> %s::vector
#                 LIMIT %s
#             """, (embedding_str, current_user["id"], request.restaurant_name, request.category, embedding_str, request.top_k))

#             results = cur.fetchall()

#             structured_results = [{
#                 "id": row[0],
#                 "filename": row[1],
#                 "chunk_id": row[2],
#                 "text_content": row[3],
#                 "similarity": float(row[4])
#             } for row in results]

#             return {
#                 "query": request.query,
#                 "restaurant": request.restaurant_name,
#                 "category": request.category,
#                 "embedding_model": request.embedding_model,
#                 "results": structured_results,
#                 "count": len(structured_results),
#                 "source": "embedding_db"
#             }

#         finally:
#             cur.close()
#             DB_POOL.putconn(conn)

#     except HTTPException:
#         raise
#     except Exception as e:
#         logger.error(f"Error querying embeddings: {str(e)}")
#         raise HTTPException(
#             status_code=500,
#             detail=f"Error querying embeddings: {str(e)}"
#         )

async def search_chat_history(request, current_user, conn):
    """Search chat history for similar questions before querying embeddings."""
    try:
        chat_history_table = "chat_history_openai" if request.embedding_model == "openai" else "chat_history_claude"

        # Generate embedding for the query
        embedding_response = openai.Embedding.create(
            input=request.query,
            model="text-embedding-ada-002"
        )
        query_embedding = embedding_response["data"][0]["embedding"]

        embedding_str = "[" + ",".join(str(x) for x in query_embedding) + "]"

        cur = conn.cursor()

        # Search in chat history using vector similarity
        cur.execute(f"""
            SELECT 
                id, question, answer, 1 - (embedding <=> %s::vector) AS similarity
            FROM {chat_history_table}
            WHERE user_id = %s 
            ORDER BY embedding <=> %s::vector
            LIMIT %s;
        """, (embedding_str, current_user["id"], embedding_str, request.top_k))

        history_results = cur.fetchall()

        structured_history_results = [{
            "id": row[0],
            "question": row[1],
            "answer": row[2],
            "similarity": float(row[3])
        } for row in history_results]

        return structured_history_results

    except Exception as e:
        logger.error(f"Error querying chat history: {str(e)}")
        return []

    # only searching for emb not history


# org

async def query_embeddings(
        request: QueryEmbeddingsRequest,
        current_user: dict = Depends(get_current_user),
        conn=Depends(get_db)
):
    """Query embeddings for similar content"""
    try:
        # Verify restaurant access
        restaurant = await verify_restaurant_access(request.restaurant_name, current_user)

        # Validate category
        # if request.category not in VALID_CATEGORIES:
        #     raise HTTPException(
        #         status_code=400,
        #         detail=f"Category must be one of: {', '.join(VALID_CATEGORIES)}"
        #     )

        # Validate embedding model
        if request.embedding_model not in ["openai", "claude"]:
            raise HTTPException(
                status_code=400,
                detail="Embedding model must be 'openai' or 'claude'"
            )

        # Generate query embedding using the appropriate model
        if request.embedding_model == "openai":
            with openai_embedding_lock:
                response = openai_client.embeddings.create(
                    input=[request.query],
                    model="text-embedding-ada-002"
                )
                query_embedding = [float(x) for x in response.data[0].embedding]
        else:  # claude
            # Try to get Claude embedding with voyage-multimodal-3
            try:
                query_embedding = get_claude_embedding(request.query)
            except Exception as e:
                logger.error(f"Error getting Claude embedding for query, falling back to OpenAI: {str(e)}")
                # Fallback to dimension-reduced OpenAI
                with openai_embedding_lock:
                    response = openai_client.embeddings.create(
                        input=[request.query],
                        model="text-embedding-ada-002"
                    )
                    full_embedding = response.data[0].embedding

                    # Reduce to 1024 dimensions
                    claude_embedding = []
                    for i in range(0, 1536, 3):
                        if i + 2 < 1536:
                            avg_value = (full_embedding[i] + full_embedding[i + 1] + full_embedding[i + 2]) / 3
                            claude_embedding.append(float(avg_value))
                        elif i + 1 < 1536:
                            avg_value = (full_embedding[i] + full_embedding[i + 1]) / 2
                            claude_embedding.append(float(avg_value))
                        else:
                            claude_embedding.append(float(full_embedding[i]))

                    # Ensure we have exactly 1024 dimensions
                    if len(claude_embedding) > 1024:
                        claude_embedding = claude_embedding[:1024]
                    elif len(claude_embedding) < 1024:
                        claude_embedding.extend([0.0] * (1024 - len(claude_embedding)))

                    # Normalize the embedding
                    embedding_array = np.array(claude_embedding)
                    norm = np.linalg.norm(embedding_array)
                    if norm > 0:
                        query_embedding = [float(x) for x in (embedding_array / norm)]
                    else:
                        query_embedding = claude_embedding

        # Set appropriate table based on requested model
        table_name = "openai_embeddings" if request.embedding_model == "openai" else "claude_embeddings"

        # Query the database for similar content using pgvector's similarity search
        conn = DB_POOL.getconn()
        try:
            cur = conn.cursor()

            # Format embedding array for PostgreSQL
            # embedding_str = "{" + ",".join(str(x) for x in query_embedding) + "}"
            embedding_str = "[" + ",".join(str(x) for x in query_embedding) + "]"

            # Use vector similarity search with cosine distance
            cur.execute(f"""
                SELECT 
                    id,
                    filename,
                    chunk_id,
                    text_content,
                    1 - (embedding <=> %s::vector) as similarity
                FROM {table_name}
                WHERE 
                    user_id = %s AND
                    restaurant_name = %s AND
                    category = %s
                ORDER BY embedding <=> %s::vector
                LIMIT %s
            """, (embedding_str, current_user["id"], request.restaurant_name, request.category, embedding_str,
                  request.top_k))

            results = cur.fetchall()

            structured_results = [{
                "id": row[0],
                "filename": row[1],
                "chunk_id": row[2],
                "text_content": row[3],
                "similarity": float(row[4])
            } for row in results]

            return {
                "query": request.query,
                "restaurant": request.restaurant_name,
                "category": request.category,
                "embedding_model": request.embedding_model,
                "results": structured_results,
                "count": len(structured_results)
            }

        finally:
            cur.close()
            DB_POOL.putconn(conn)
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error querying embeddings: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Error querying embeddings: {str(e)}"
        )


###################################################################

def create_chunks_from_docx(file_content, max_chunk_length=1000):
    doc = Document(io.BytesIO(file_content))
    text = [para.text for para in doc.paragraphs if para.text.strip()]

    # Combine all paragraphs into a single string
    combined_text = " ".join(text)

    # Split the combined text into chunks of max_chunk_length characters
    chunks = []
    for i in range(0, len(combined_text), max_chunk_length):
        chunks.append(combined_text[i:i + max_chunk_length])

    return chunks


# Function to extract text from different document types
def extract_text_from_doc(file_content: bytes, doc_type: str):
    # if doc_type == "csv":
    #     df = pd.read_csv(io.StringIO(file_content.decode("utf-8")))
    #     return [" ".join([f"{col}: {row[col]}" for col in df.columns]) for _, row in df.iterrows()]
    if doc_type == "csv":
        df = pd.read_csv(io.StringIO(file_content.decode("utf-8")))

        # Create chunks of 5 rows each
        chunk_size = 20
        chunks = []

        # Iterate over the dataframe in chunks of 5 rows
        for i in range(0, len(df), chunk_size):
            chunk = df.iloc[i:i + chunk_size]
            chunk_str = "\n".join(
                [" ".join([f"{col}: {row[col]}" for col in chunk.columns]) for _, row in chunk.iterrows()])
            chunks.append(chunk_str)

        return chunks


    elif doc_type == "pdf":
        text_chunks = []
        with pdfplumber.open(io.BytesIO(file_content)) as pdf:
            for page in pdf.pages:
                text_chunks.append(page.extract_text())
        return text_chunks

    elif doc_type == "docx":
        chunks = create_chunks_from_docx(file_content, max_chunk_length=1500)
        # doc = Document(io.BytesIO(file_content))
        # return [para.text for para in doc.paragraphs if para.text.strip()]
        return chunks

    elif doc_type in ["xls", "xlsx"]:
        wb = openpyxl.load_workbook(io.BytesIO(file_content))
        text_chunks = []
        for sheet in wb.sheetnames:
            ws = wb[sheet]
            for row in ws.iter_rows(values_only=True):
                text_chunks.append(" ".join(map(str, row)))
        return text_chunks

    else:
        raise ValueError("Unsupported document type")


async def generate_embeddings_all_docs(
        s3_key: str,
        restaurant_name: str,
        category: str,
        filename: str,
        embedding_model: str,
        current_user: dict,
        batch_size: int = 10
) -> Dict[str, Any]:
    """Generate embeddings for different document formats stored in S3"""
    try:
        # Detect document type
        doc_type = filename.split(".")[-1].lower()

        # Retrieve file from S3
        response = s3_client.get_object(Bucket=BUCKET_NAME, Key=s3_key)
        file_content = response['Body'].read()

        # Extract text based on document type
        chunks = extract_text_from_doc(file_content, doc_type)

        # Prepare chunk data
        processed_chunks = []
        for index, text in enumerate(chunks):
            chunk_id = index  # Can be modified for hashing if needed
            processed_chunks.append((text, current_user["id"], restaurant_name, category, filename, chunk_id))

        # Choose embedding model processing function  openai
        if embedding_model == "openai":
            process_func = process_chunk_openai
            table_name = "openai_embeddings"
        else:
            process_func = process_chunk_claude
            table_name = "claude_embeddings"

        # Process chunks in parallel
        all_embeddings = []

        with ThreadPoolExecutor(max_workers=CPU_COUNT) as executor:
            futures = [executor.submit(process_func, *chunk) for chunk in processed_chunks]
            for future in concurrent.futures.as_completed(futures):
                result = future.result()
                if result:
                    all_embeddings.append(result)

                if len(all_embeddings) >= batch_size:
                    conn = DB_POOL.getconn()
                    try:
                        cur = conn.cursor()
                        execute_values(cur,
                                       f"INSERT INTO {table_name} (user_id, restaurant_name, category, filename, chunk_id, text_content, embedding) VALUES %s",
                                       all_embeddings
                                       )
                        conn.commit()
                        logger.info(
                            f"✔ Inserted {len(all_embeddings)} embeddings into {table_name} with multithreading ...")
                        all_embeddings = []

                    finally:
                        cur.close()
                        DB_POOL.putconn(conn)

        # Insert any remaining embeddings
        if all_embeddings:
            conn = DB_POOL.getconn()
            try:
                cur = conn.cursor()
                execute_values(cur,
                               f"INSERT INTO {table_name} (user_id, restaurant_name, category, filename, chunk_id, text_content, embedding) VALUES %s",
                               all_embeddings
                               )
                conn.commit()
                logger.info(f"🚩 Inserted {len(all_embeddings)} embeddings into {table_name} without multithreading ...")
            finally:
                cur.close()
                DB_POOL.putconn(conn)

        return {
            "message": f"{embedding_model.capitalize()} embeddings for {filename} processed and stored successfully.",
            "restaurant": restaurant_name,
            "category": category,
            "chunks_processed": len(chunks),
            "embedding_model": embedding_model
        }

    except Exception as e:
        logger.error(f"Error generating embeddings: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


######################################################################

def save_chat_history_db(
        user_id: str,
        embedding_model: str,
        restaurant_name: str,
        category: str,
        filename: str,
        question: str,
        answer: str,
):
    """Save chat history embeddings in the database with error handling."""
    try:
        # Validate embedding model
        if embedding_model not in ["openai", "claude"]:
            logger.error(f"Invalid embedding model: {embedding_model}")
            return

        # Prepare text chunk
        chunks = [f"Question: {question}, Answer: {answer}"]
        processed_chunks = [
            (text, user_id, restaurant_name, category, filename, idx)
            for idx, text in enumerate(chunks)
        ]

        # Select processing function and table
        process_func = process_chunk_openai if embedding_model == "openai" else process_chunk_claude
        table_name = "chat_history_openai" if embedding_model == "openai" else "chat_history_claude"

        # Process chunks in parallel
        all_embeddings = []
        with ThreadPoolExecutor(max_workers=5) as executor:
            futures = {executor.submit(process_func, *chunk): chunk for chunk in processed_chunks}

            for future in concurrent.futures.as_completed(futures):
                try:
                    result = future.result()
                    if result:
                        all_embeddings.append(result)
                except Exception as e:
                    logger.error(f"Error processing chunk {futures[future]}: {str(e)}")

        # Prepare embeddings for insertion
        if not all_embeddings:
            logger.warning("No embeddings generated, skipping database insertion.")
            return

        formatted_embeddings = [
            (user_id, restaurant_name, category, filename, chunk_id, text_content, embedding, embedding_model)
            for user_id, restaurant_name, category, filename, chunk_id, text_content, embedding in all_embeddings
        ]

        # Insert embeddings into database
        conn = DB_POOL.getconn()
        try:
            cur = conn.cursor()
            execute_values(
                cur,
                f"INSERT INTO {table_name} (user_id, restaurant_name, category, filename, chunk_id, text_content, embedding, model) VALUES %s",
                formatted_embeddings
            )
            conn.commit()
        except Exception as e:
            logger.error(f"Database insertion error: {str(e)}")
        finally:
            cur.close()
            DB_POOL.putconn(conn)

    except Exception as e:
        logger.error(f"Unexpected error in save_chat_history_db: {str(e)}")

    # Function to initialize everything


def init_rag_module():
    """Initialize the RAG module"""
    try:
        init_embedding_tables()
        logger.info("RAG module initialized successfully")
    except Exception as e:
        logger.error(f"Error initializing RAG module: {str(e)}")
        raise


# Function to initialize everything
def init_directories():
    """Initialize the RAG module"""
    try:
        # If it doesn't exist, create it
        if not os.path.exists("uploads/profile_images/"):
            os.makedirs("uploads/profile_images")
            logger.info("uploads/profile_images created successfully")
        else:
            logger.info("uploads/profile_images initialized successfully")

    except Exception as e:
        logger.error(f"Error initializing DIR module: {str(e)}")
        raise


# Make sure to create tables on application startup
init_rag_module()
init_directories()
