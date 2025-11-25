from fastapi import FastAPI, File, UploadFile, HTTPException, WebSocket, WebSocketDisconnect, Form
from typing import List, Optional, Dict, Any
from pydantic import BaseModel
import anthropic
import csv
import os
import io
import PyPDF2
from pdfminer.high_level import extract_text
from pdfminer.layout import LAParams
import json
from datetime import datetime
import logging
from jwt import PyJWTError
from docx import Document
import uuid
import pandas as pd


import json
import os

from datetime import datetime


# Set up a logger


import boto3


from botocore.exceptions import ClientError

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Models

from src.config import notification_manager
# from src.chat_gpt import notification_manager

logger.info(f"notification_manager id: {id(notification_manager)}")

class ChatRequest(BaseModel):
    question: str
    file_type: Optional[str] = "all"
    file_index: Optional[int] = -1
    include_history: bool = True

# File Content Class
class FileContent:
    def __init__(self, filename: str, content_type: str):
        self.filename = filename
        self.content_type = content_type  # "pdf" or "csv"
        self.text_content = ""
        self.structured_content = {}
        self.summary = ""


class TranscriptChatRequest(BaseModel):
    file_type: str = "all"
    file_index: int = -1
    include_history: bool = True

# Chat History Class

# Configure logging


# AWS S3 Configuration
# Chat History Class
class ChatHistory:
    def __init__(self):
        self.conversations = []
        self.max_history = 10
        self.history_file = "chat_history.json"
        self.load_from_file()

    def add_conversation(self, query: dict, response: str, user_email: str):
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        conversation_id = str(uuid.uuid4())
        
        self.conversations.append({
            "id": conversation_id,
            "timestamp": timestamp,
            "query": query,
            "response": response,
            "user_email": user_email,
            "file_types_present": query.get('file_type', 'all'),
            "model" : "c"
        })
        
        if len(self.conversations) > self.max_history:
            self.conversations = self.conversations[-self.max_history:]
        
        self.save_to_file()

    def delete_conversation_by_id(self, conversation_id: str, user_email: str) -> bool:
        """Delete a specific conversation by its ID for a user"""
        try:
            # Find the conversation with matching ID
            filtered_conversations = [
                conv for conv in self.conversations 
                if conv.get('user_email') == user_email
            ]
            
            for i, conv in enumerate(filtered_conversations):
                if conv.get('id') == conversation_id:
                    self.conversations.remove(conv)
                    self.save_to_file()
                    return True
            return False
        except Exception as e:
            logger.error(f"Error deleting conversation: {str(e)}")
            return False


    def get_context(self, user_email: str) -> str:
        if not self.conversations:
            return ""
        
        context = "\nPrevious conversations:\n"
        for conv in self.conversations:
            if conv.get('user_email') == user_email and conv.get('model') == "c":  # Use get() for safe access
                context += f"\nTime: {conv['timestamp']}\n"
                context += f"Q: {conv['query']['question']}\n"
                context += f"A: {conv['response']}\n"
                context += "-" * 40 + "\n"
        return context

    def save_to_file(self):
        try:
            with open(self.history_file, 'w', encoding='utf-8') as f:
                json.dump(self.conversations, f, indent=2, ensure_ascii=False)
        except Exception as e:
            print(f"Error saving history: {str(e)}")

    def load_from_file(self):
        try:
            if os.path.exists(self.history_file):
                with open(self.history_file, 'r', encoding='utf-8') as f:
                    loaded_conversations = json.load(f)
                    self.conversations=[
                        conv for conv in loaded_conversations
                        if isinstance(conv, dict) and 'user_email' in conv 
                    ]
            else:
                logger.info("History file does not exist. CXXXXXXXXXXXXXXXXXXX")
                self.conversations = []

                    # adding validation
        except Exception as e:
            logger.error(f"Error loading history: {str(e)}")
            self.conversations = []

    def clear_history(self,user_email: Optional[str]=None):
        if user_email:
            self.conversations = [
                conv for conv in self.conversations
                if conv.get('user_email') != user_email
            ]
        else:
            self.conversations = []
        self.save_to_file()

FOHBOH_SYSTEM_PROMPT = """You are Fohboh, a knowledgeable Restaurant Assistant. Your role is to help restaurant owners and managers with their questions and tasks.

Key responsibilities:
Provide expert advice on restaurant operations and management
Analyze restaurant data and reports
Offer menu suggestions and optimization
Help with staffing and scheduling
Guide on health and safety regulations
Assist with customer service strategies
Provide marketing and promotion ideas
Help with inventory management
Guide on cost control and financial planning
Share industry best practices and trends

Approach:
Always maintain a professional and supportive tone
Provide practical, actionable advice
Use industry-specific terminology appropriately
Focus on solutions that are realistic for restaurants
Consider both operational and financial impacts
Reference relevant regulations and standards when applicable
Provide specific examples and scenarios when helpful
When responding to questions requiring calculations or data analysis, assume typical or standard data parameters if explicit data is not provided, and respond with a valid, logical answer based on those assumptions. Avoid stating that data is missing or unavailable.

Important Instruction:
When addressing labor cost calculations, assume there are 3 employees who have each worked 6 hours. Use this as the foundational example to illustrate how labor cost is calculated unless otherwise specified.

Note:

- Always aim to respond using the provided data and avoid requesting additional information each time.
- Please do not include any random data or calculations that are not provided.
- If you receive an irrelevant question (not related to fohboh), respond with: 'Sorry, I am Fohbot, an AI companion designed to assist restaurant owners and managers with restaurant-related tasks and inquiries—such as operations, staffing, menus, customer service, and marketing. I'm unable to help with anything outside of this scope.'
- Always aim to respond using the provided or assumed data and avoid requesting additional information each time.
- Avoid any disclaimers or phrases such as “assume,” “if,” “might,” “could,” or similar.
- Maintain a confident, factual tone throughout every answer.
- Provide concise, to-the-point answers.
- Avoid long explanations or essays.
- Deliver instant, actionable responses that can be used right away.
- Keep answers short but complete enough to be useful.

- Whenever a comparison is involved, always present the comparison in a clear table format. Do not provide comparisons as plain text or lists for comparison.

"""






# File Processor Class
class FileProcessor:
    def __init__(self):
        self.api_key = "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"
        self.client = anthropic.Anthropic(
            api_key=self.api_key
            
        )
        self.system_prompt = FOHBOH_SYSTEM_PROMPT
        self.MAX_FILES = 3
        self.files = {}
        self.history = ChatHistory()
        
    def extract_text_with_layout(self, pdf_file: io.BytesIO) -> str:
        try:
            laparams = LAParams(
                line_margin=0.5,
                word_margin=0.1,
                char_margin=2.0,
                boxes_flow=0.5,
                detect_vertical=True
            )
            return extract_text(pdf_file, laparams=laparams)
        except Exception as e:
            logger.error(f"Error in text extraction: {str(e)}")
            return ""

    def process_pdf(self, pdf_file: io.BytesIO, filename: str, user_email: str) -> bool:
        try:
            file_content = FileContent(filename, "pdf")
            text_content = self.extract_text_with_layout(pdf_file)
            file_content.text_content = text_content

            pdf_file.seek(0)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            file_content.structured_content = {
                "num_pages": len(pdf_reader.pages),
                "metadata": pdf_reader.metadata if pdf_reader.metadata else {},
                "content_preview": text_content[:1000] if text_content else "No text content available"
            }

            file_content.summary = f"PDF File: {filename}\nPages: {len(pdf_reader.pages)}\nPreview: {text_content[:200]}..."
            
            if user_email not in self.files:
                self.files[user_email] = []
            self.files[user_email].append(file_content)
            return True
        except Exception as e:
            logger.error(f"Error processing PDF {filename}: {str(e)}")
            return False

    def process_csv(self, csv_file: io.BytesIO, filename: str, user_email: str) -> bool:
        try:
            file_content = FileContent(filename, "csv")
            text = csv_file.read().decode('utf-8')
            csvfile = io.StringIO(text)
            reader = csv.reader(csvfile)
            rows = list(reader)
            if not rows:
                return False

            headers = rows[0]
            data = rows[1:]

            file_content.text_content = text
            file_content.structured_content = {
                "headers": headers,
                "total_rows": len(data),
                "sample_data": data[:5]
            }

            sample_data_str = "\n".join([", ".join(row) for row in data[:3]])
            file_content.summary = f"CSV File: {filename}\nColumns: {', '.join(headers)}\nTotal Rows: {len(data)}\nSample Data:\n{sample_data_str}"

            if user_email not in self.files:
                self.files[user_email] = []
            self.files[user_email].append(file_content)
            return True
        except Exception as e:
            logger.error(f"Error processing CSV {filename}: {str(e)}")
            return False

    def process_xlsx(self, xlsx_file: io.BytesIO, filename: str, user_email: str) -> bool:
        try:
            file_content = FileContent(filename, "xlsx")

            # Read the Excel file
            xlsx_file.seek(0)
            df = pd.read_excel(xlsx_file, engine="openpyxl")

            if df.empty:
                return False

            headers = df.columns.tolist()
            data = df.head(5).values.tolist()

            file_content.text_content = df.to_csv(index=False)  # Store as CSV text
            file_content.structured_content = {
                "headers": headers,
                "total_rows": len(df),
                "sample_data": data
            }

            sample_data_str = "\n".join([", ".join(map(str, row)) for row in data])
            file_content.summary = f"XLSX File: {filename}\nColumns: {', '.join(headers)}\nTotal Rows: {len(df)}\nSample Data:\n{sample_data_str}"

            if user_email not in self.files:
                self.files[user_email] = []
            self.files[user_email].append(file_content)
            return True
        except Exception as e:
            logger.error(f"Error processing XLSX {filename}: {str(e)}")
            return False

    def process_docx(self, docx_file: io.BytesIO, filename: str, user_email: str) -> bool:
        try:
            file_content = FileContent(filename, "docx")

            docx_file.seek(0)
            doc = Document(docx_file)

            # Extract text from paragraphs
            text_content = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])

            file_content.text_content = text_content
            file_content.structured_content = {
                "num_paragraphs": len(doc.paragraphs),
                "content_preview": text_content[:1000] if text_content else "No text content available"
            }

            file_content.summary = f"DOCX File: {filename}\nParagraphs: {len(doc.paragraphs)}\nPreview: {text_content[:200]}..."

            if user_email not in self.files:
                self.files[user_email] = []
            self.files[user_email].append(file_content)
            return True
        except Exception as e:
            logger.error(f"Error processing DOCX {filename}: {str(e)}")
            return False


    def clear_files(self, user_email: str):
        if user_email in self.files:
            del self.files[user_email]
        logger.info(f"Files cleared for user: {user_email}")

    async def process_files(self, files: List[UploadFile], user_email: str) -> dict:
        if len(files) > self.MAX_FILES:
            raise HTTPException(
                status_code=400, 
                detail=f"Maximum {self.MAX_FILES} files allowed. You tried to upload {len(files)} files."
            )

        self.clear_files(user_email)
        
        processed_files = {
            "pdfs": 0,
            "csvs": 0,
            "docxs":0,
            "xlsxs":0,
            "errors": []
        }

        for file in files:
            ext = os.path.splitext(file.filename)[1].lower()
            try:
                file_bytes = io.BytesIO(await file.read())
                
                if ext == '.pdf':
                    if self.process_pdf(file_bytes, file.filename, user_email):
                        processed_files["pdfs"] += 1
                    else:
                        processed_files["errors"].append(f"Failed to process PDF: {file.filename}")
                
                elif ext == '.csv':
                    if self.process_csv(file_bytes, file.filename, user_email):
                        processed_files["csvs"] += 1
                    else:
                        processed_files["errors"].append(f"Failed to process CSV: {file.filename}")

                elif ext == '.docx':
                    if self.process_docx(file_bytes, file.filename, user_email):
                        processed_files["docxs"] += 1
                    else:
                        processed_files["errors"].append(f"Failed to process docx: {file.filename}")
                elif ext == '.xlsx':
                    if self.process_xlsx(file_bytes, file.filename, user_email):
                        processed_files["xlsxs"] += 1
                    else:
                        processed_files["errors"].append(f"Failed to process xlsx: {file.filename}")
                
                else:
                    processed_files["errors"].append(f"Unsupported file type: {ext}")
                    
            except Exception as e:
                processed_files["errors"].append(f"Error processing {file.filename}: {str(e)}")
                continue

        return processed_files
    
    async def chat_from_transcript(self,conn, current_user, transcript_text: str, request: TranscriptChatRequest, user_email: str) -> str:
        """Process a transcript and generate a response."""
        try:
            response = await self.generate_response(
                question=transcript_text,
                user_email=user_email,
                file_type=request.file_type,
                file_index=request.file_index,
                include_history=request.include_history
            )
            
            cur = conn.cursor()
      

            return {
                "transcript": transcript_text,
                "response": response,
                "has_history": len(self.history.conversations) > 0
            }
        except Exception as e:
            logger.error(f"Error in chat_from_transcript: {str(e)}")
            raise HTTPException(
                status_code=500,
                detail=f"Error processing transcript: {str(e)}"
            )
        

    async def generate_streaming_response(
        self,
        question: str,
        client_id: str,
        manager: 'ConnectionManager',
        user_email: str,
        file_type: str = "all",
        file_index: int = -1,
        include_history: bool = True
        ) -> None:
            try:
                messages = []
                messages.append({
                    "role": "assistant",
                    "content": self.system_prompt
                })
                
                if include_history:
                    history_context = self.history.get_context(user_email)
                    if history_context:
                        messages.append({
                            "role": "assistant",
                            "content": history_context
                        })
                
                if user_email in self.files and self.files[user_email]:
                    if file_type == "all":
                        context_parts = []
                        for file in self.files[user_email]:
                            context_parts.append(f"\n{'='*50}\n{file.summary}\n{'='*50}")
                        context = "\n".join(context_parts)
                    else:
                        matching_files = [f for f in self.files[user_email] if f.content_type == file_type]
                        if not matching_files:
                            await manager.send_message(f"No {file_type} files found.", client_id)
                            return
                        
                        if file_index >= 0 and file_index < len(matching_files):
                            file = matching_files[file_index]
                            context = f"File content from {file.filename}:\n{file.text_content}"
                        else:
                            await manager.send_message(f"Invalid file index {file_index} for {file_type} files.", client_id)
                            return
        
                    messages.append({
                        "role": "user",
                        "content": f"Context:\n{context}\n\nQuestion: {question}"
                    })
                else:
                    messages.append({
                        "role": "user",
                        "content": question
                    })
        
                # Stream response
                with self.client.messages.stream(
                    #model="claude-3-opus-20240229",
                    model="claude-sonnet-4-5-20250929",
                    max_tokens=1024,
                    temperature=0.7,
                    messages=messages
                    
                ) as stream:
                    full_response = ""
                    for chunk in stream:
                        if chunk.type == 'content_block_delta':
                            text_chunk = chunk.delta.text
                            if text_chunk:
                                await manager.send_message(text_chunk, client_id)
                                full_response += text_chunk
        
                # Save to history after complete response
                self.history.add_conversation(
                    {"question": question, "file_type": file_type, "file_index": file_index},
                    full_response,
                    user_email
                )
        
            except Exception as e:
                logger.error(f"Streaming error: {str(e)}")
                await manager.send_message(f"Error: {str(e)}", client_id)
    

    async def generate_response_with_embeddings(
        self,
        question: str,
        user_email: str,
        user_id:int,
        embedding_info: Optional[Dict[str, Any]] = None,
        file_type: str = "all",
        file_index: int = -1,
        include_history: bool = True
    ) -> str:
        """Generate response using Claude API with optional embedding augmentation."""
        try:
            # Removed predefined answers - now always generates responses from actual data
            
            messages = []
            messages.append({
                "role": "assistant",
                "content": self.system_prompt
            })
            
            # Extract personal information from chat history
            user_info = {}
            if include_history:
                # First, gather user information from previous interactions
                user_history = [
                    conv for conv in self.history.conversations 
                    if conv.get('user_email') == user_email
                ]
                
                # Look for personal information in past conversations
                for conv in user_history:
                    # Check for name introduction patterns in questions
                    q = conv.get('query', {}).get('question', '').lower()
                    if "my name is" in q:
                        # Extract name after "my name is"
                        name_part = q.split("my name is")[1].strip()
                        # Take first word as name
                        name = name_part.split()[0] if ' ' in name_part else name_part
                        user_info['name'] = name
                
                # Now add the regular history context
                history_context = self.history.get_context(user_email)
                if history_context:
                    messages.append({
                        "role": "assistant",
                        "content": history_context
                    })
            
            # Add embedding context if available
            embedding_context = ""
            if embedding_info:
                # Handle multi-restaurant format
                if embedding_info.get("multi_restaurant", False) and embedding_info.get("restaurant_results"):
                    embedding_context += f"\n\nRelevant information from {embedding_info['restaurant_count']} restaurants:\n"
                    
                    for restaurant_data in embedding_info["restaurant_results"]:
                        restaurant_name = restaurant_data["restaurant_name"]
                        embedding_context += f"\n\n=== Restaurant: {restaurant_name} ===\n"
                        
                        for category_data in restaurant_data["results"]:
                            category = category_data["category"]
                            embedding_context += f"\n--- {category} Category ---\n"
                            
                            for file_data in category_data["files"]:
                                filename = file_data["filename"]
                                embedding_context += f"\nFrom file '{filename}':\n"
                                
                                # Include the top results from each file
                                for idx, result in enumerate(file_data["results"]):
                                    embedding_context += f"- {result['text_content']}\n"
                                    if idx >= 2:  # Limit to 3 results per file to avoid making the context too large
                                        break
                # Handle single restaurant format (original)
                elif embedding_info.get("results"):
                    restaurant_name = embedding_info['restaurant_name']
                    embedding_context += f"\n\nRelevant information from {restaurant_name} restaurant:\n"
                    
                    for category_data in embedding_info["results"]:
                        category = category_data["category"]
                        embedding_context += f"\n--- {category} Category ---\n"
                        
                        for file_data in category_data["files"]:
                            filename = file_data["filename"]
                            embedding_context += f"\nFrom file '{filename}':\n"
                            
                            # Include the top results from each file
                            for idx, result in enumerate(file_data["results"]):
                                embedding_context += f"- {result['text_content']}\n"
                                if idx >= 2:  # Limit to 3 results per file to avoid making the context too large
                                    break
            
            # Check if there are files to analyze
            if user_email in self.files and self.files[user_email]:
                # File-based conversation
                if file_type == "all":
                    context_parts = []
                    for file in self.files[user_email]:
                        context_parts.append(f"\n{'='*50}\n{file.summary}\n{'='*50}")
                    context = "\n".join(context_parts)
                else:
                    matching_files = [f for f in self.files[user_email] if f.content_type == file_type]
                    if not matching_files:
                        return f"No {file_type} files found."
                    
                    if file_index >= 0 and file_index < len(matching_files):
                        file = matching_files[file_index]
                        context = f"File content from {file.filename}:\n{file.text_content}"
                    else:
                        return f"Invalid file index {file_index} for {file_type} files."
                
                # Add the file context, embedding context, and question together
                full_context = context
                if embedding_context:
                    full_context += "\n\n" + embedding_context
                    
                messages.append({
                    "role": "user",
                    "content": f"Context:\n{full_context}\n\nQuestion: {question}"
                })
            else:
                # General conversation mode
                if embedding_context:
                    messages.append({
                        "role": "user",
                        "content": f"Context:\n{embedding_context}\n\nQuestion: {question}"
                    })
                else:
                    messages.append({
                        "role": "user",
                        "content": question
                    })
            
            # Add specific instructions for formatting the response
            instructions = """Please format your response as follows:
                            1. Provide a clear, concise answer to the question using the available data
                            2. Include specific numeric values when applicable
                            3. Always end with 2 relevant follow-up questions labeled as "Follow-up questions:"
                            4. The follow-up questions should help the user gain deeper insights or explore related areas
                            5. Do NOT include data source citations at the end of your response
                            6. Make your follow-up questions insightful and business-relevant"""

            # Add user information to the instructions if available
            if user_info:
                instructions += "\n\nImportant user information from previous conversations:"
                if 'name' in user_info:
                    instructions += f"\n- User's name is {user_info['name']}"
                
            messages.append({
                "role": "user",
                "content": instructions
            })
            
            try:
                
                with self.client.messages.stream(
                    #model="claude-3-opus-20240229",
                    model="claude-sonnet-4-5-20250929",
                    max_tokens=1024,
                    messages=messages
                    
                ) as stream:
                    response_text = ""
                    for chunk in stream:
                        if chunk.type == 'content_block_delta':
                            text_chunk = chunk.delta.text
                            if text_chunk:
                                response_text += text_chunk
                                try:
                            # Attempt to broadcast
                                    message={
                                            "type": "chat",
                                            "action": "token",
                                            "data": {
                                                "token": text_chunk,
                                                "timestamp": datetime.now().isoformat()
                                            }
                                        }
                                    await notification_manager.publish_notification(user_id, message)
                                    
                                    # await notification_manager.broadcast_to_user(
                                    #     user_id=user_id,
                                    #     message=message)
                                    # logger.info(f"Successfully broadcast token to user {user_id}")
                                    
                                except Exception as e:
                                    logger.error(f"Error broadcasting chat token: {str(e)}", exc_info=True)        
                                
                                
                                
                # Generate response with Claude
                # response = self.client.messages.create(
                #     model="claude-3-opus-20240229",
                #     max_tokens= 1024,
                #     messages=messages
                # )
                
                # response_text = response.content[0].text
                
                # Remove "FohBoh Response:" prefix if it exists
                if response_text.strip().startswith("FohBoh Response:"):
                    response_text = response_text.replace("FohBoh Response:", "", 1).strip()
                
                # Save to history
                self.history.add_conversation(
                    {"question": question, "file_type": file_type, "file_index": file_index},
                    response_text,
                    user_email
                )
                
                return response_text
                
            except Exception as e:
                logger.error(f"API Error: {str(e)}")
                raise HTTPException(status_code=500, detail=f"Error in API call: {str(e)}")
        
        except Exception as e:
            # return f"An error occurred ...: {str(e)}"
            raise HTTPException(status_code=500, detail=f"An error occurred ...: {str(e)}")


            

        
    async def generate_response(
    self, 
    question: str, 
    user_email: str,
    file_type: str = "all", 
    file_index: int = -1, 
    include_history: bool = True
    ) -> str:
        try:
            messages = []
            messages.append({
                "role": "assistant",
                "content": self.system_prompt
            })
            
            # Add chat history if requested
            if include_history:
                history_context = self.history.get_context(user_email)
                if history_context:
                    messages.append({
                        "role": "assistant",
                        "content": history_context
                    })
            
            # Check if there are files to analyze
            if user_email in self.files and self.files[user_email]:
                # File-based conversation
                if file_type == "all":
                    context_parts = []
                    for file in self.files[user_email]:
                        context_parts.append(f"\n{'='*50}\n{file.summary}\n{'='*50}")
                    context = "\n".join(context_parts)
                else:
                    matching_files = [f for f in self.files[user_email] if f.content_type == file_type]
                    if not matching_files:
                        return f"No {file_type} files found."
                    
                    if file_index >= 0 and file_index < len(matching_files):
                        file = matching_files[file_index]
                        context = f"File content from {file.filename}:\n{file.text_content}"
                    else:
                        return f"Invalid file index {file_index} for {file_type} files."
                
                messages.append({
                    "role": "user",
                    "content": f"Context:\n{context}\n\nQuestion: {question}"
                })
            else:
                # General conversation mode
                messages.append({
                    "role": "user",
                    "content": question
                })
            
            try:
               
                response = self.client.messages.create(
                   # model="claude-3-opus-20240229",
                    model="claude-sonnet-4-5-20250929",
                    max_tokens=1024,
                    messages=messages
                )
                
                response_text = response.content[0].text
                
                self.history.add_conversation(
                    {"question": question, "file_type": file_type, "file_index": file_index},
                    response_text,
                    user_email
                )
                
                return response_text
                
            except Exception as e:
                logger.error(f"API Error: {str(e)}")
                raise HTTPException(status_code=500, detail=f"Error in API call: {str(e)}")
        
        except Exception as e:
            return f"An error occurred: {str(e)}"

    
def _get_updated_prompt(prompt_n: int) -> str:
    # Define the use case scenarios in detail
    USE_CASE_1 = """
        Use Case: Retrieve sales data from the POS system for last week across all locations and match it with food cost data from the inventory system. 
        Generate a report showing the top-selling menu items by revenue, and calculate their profitability (revenue - food cost). 
        Provide a list of the most profitable items, highlighting both their sales figures and food cost.
    """
    
    USE_CASE_2 = """
        Use Case: Access labor data from the workforce management system and POS system. 
        Forecast the labor costs for this week by daypart (morning, lunch, dinner) and compare them with last week's actual labor costs by daypart. 
        Identify any trends or discrepancies and provide recommendations for staffing adjustments to improve cost efficiency.
    """
    
    USE_CASE_3 = """
        Use Case: Pull historical sales data from Toast for the past 30 days and combine it with current inventory levels from the inventory system. 
        Based on sales trends, predict how much inventory is needed for the top-selling items for the upcoming week. 
        Suggest appropriate order quantities to avoid over-ordering and stockouts.
    """
    
    USE_CASE_4 = """
        Use Case: Retrieve labor data and sales data from Toast for dinner shifts over the past month. 
        Identify the employees who have the highest sales per labor hour during dinner shifts. 
        Provide insights on how to optimize staffing and scheduling based on the top performers.
    """
    
    USE_CASE_5 = """
        Use Case: Analyze menu performance data from Toast for the lunch period and calculate the food cost for each menu item using the inventory system data. 
        Identify the least profitable items during lunch hours and recommend whether these should be removed, repriced, or reworked to improve profitability.
    """
    
    USE_CASE_6 = """
        Use Case: Analyze historical labor data and sales data from Toast for breakfast service over the past month. 
        Based on this data, calculate the optimal number of servers required to meet the target labor cost percentage while ensuring efficient service. 
        Provide staffing recommendations to meet these targets.
    """
    
    USE_CASE_7 = """
        Use Case: Retrieve inventory data from the inventory system and sales data from Toast. 
        Calculate the shrinkage (variance between actual and theoretical inventory usage) for last week. 
        Identify the main contributors to shrinkage (e.g., waste, theft, over-portioning) and provide actionable recommendations for minimizing losses.
    """
    
    USE_CASE_8 = """
        Use Case: Analyze historical sales data from Toast and labor costs from the workforce management system for past Friday dinner services. 
        Based on these trends, recommend an optimized staffing schedule for next Friday's dinner service to ensure adequate staffing without overspending on labor. 
        Provide specific staffing numbers and shift timings.
    """
    
    USE_CASE_9 = """
        Use Case: Access menu item data from Toast and food cost data from the inventory system. 
        Identify the menu items with the highest variance between expected and actual food costs. 
        Provide recommendations on how to reduce this variance, such as improving portion control or adjusting pricing strategies.
    """
    
    USE_CASE_10 = """
        Use Case: Analyze POS data to calculate the customer flow and average ticket size for last weekend. 
        Compare this with the number of staff scheduled for that period to assess staffing efficiency. 
        Provide insights into whether the staffing levels were appropriate and suggest adjustments to the staff schedule based on the observed customer flow and ticket size patterns.
    """
    
    use_cases = {
        1: USE_CASE_1,
        2: USE_CASE_2,
        3: USE_CASE_3,
        4: USE_CASE_4,
        5: USE_CASE_5,
        6: USE_CASE_6,
        7: USE_CASE_7,
        8: USE_CASE_8,
        9: USE_CASE_9,
        10: USE_CASE_10
    }
    
    
    # Get the appropriate use case based on the prompt number
    updated_prompt = f"{FOHBOH_SYSTEM_PROMPT}\n{use_cases.get(prompt_n)}"
    
    return updated_prompt





# WebSocket connection manager
class ConnectionManager:
    def __init__(self):
        self.active_connections: Dict[str, WebSocket] = {}

    async def connect(self, websocket: WebSocket, client_id: str):
        await websocket.accept()  # Accept the connection here
        self.active_connections[client_id] = websocket

    def disconnect(self, client_id: str):
        self.active_connections.pop(client_id, None)

    async def send_message(self, message: str, client_id: str):
        if client_id in self.active_connections:
            await self.active_connections[client_id].send_text(message)

# Initialize global instances
processor = FileProcessor()
manager = ConnectionManager()
